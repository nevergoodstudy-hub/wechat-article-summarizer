"""Word文档导出器 - 使用 html2docx 实现高保真转换

支持生成Word 2013+现代格式文档，避免在Microsoft 365中显示为兼容性模式。
通过设置compatibilityMode = 15，确保文档使用Word 2013+的现代XML结构。
"""

import base64
import io
import re
from pathlib import Path
from typing import TYPE_CHECKING, Any
from urllib.parse import urljoin

from bs4 import BeautifulSoup
from loguru import logger

from ....domain.entities import Article
from ....shared.exceptions import ExporterError
from ....shared.utils.ssrf_protection import safe_fetch_sync
from .base import BaseExporter

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
else:
    DocxDocument = Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from lxml import etree

    _docx_available = True
except ImportError:
    _docx_available = False

try:
    from html2docx import html2docx

    _html2docx_available = True
except ImportError:
    _html2docx_available = False


# Word兼容性模式值
# 11 = Word 2003, 12 = Word 2007, 14 = Word 2010, 15 = Word 2013+
WORD_COMPATIBILITY_MODE_2013 = "15"
MAX_REMOTE_IMAGE_BYTES = 10 * 1024 * 1024


class WordExporter(BaseExporter):
    """
    Word文档导出器

    使用 html2docx 将文章的 HTML 内容高保真转换为 Word 文档，
    保留原始排版、图片位置和文字格式。
    """

    def __init__(self, output_dir: str = "./output"):
        self._output_dir = Path(output_dir)

    @property
    def name(self) -> str:
        return "word"

    @property
    def target(self) -> str:
        return "word"

    def is_available(self) -> bool:
        """检查依赖是否可用"""
        return _docx_available

    def export(
        self,
        article: Article,
        path: str | None = None,
        **options,
    ) -> str:
        """导出为Word文档"""
        if not _docx_available:
            raise ExporterError("Word导出需要安装 python-docx: pip install python-docx")

        # 确定输出路径
        if path:
            output_path = Path(path)
            if output_path.is_dir():
                output_path = output_path / self._generate_filename(article)
        else:
            self._output_dir.mkdir(parents=True, exist_ok=True)
            output_path = self._output_dir / self._generate_filename(article)

        # 生成Word文档
        try:
            # 直接使用手动构建模式，更稳定
            doc = self._generate_document_manual(article, **options)
            doc.save(str(output_path))

            logger.info(f"Word导出成功: {output_path}")
            return str(output_path)
        except Exception as e:
            raise ExporterError(f"Word导出失败: {e}") from e

    def _generate_filename(self, article: Article) -> str:
        """生成文件名"""
        safe_title = re.sub(r'[\\/*?:"<>|]', "", article.title)
        safe_title = safe_title[:50]
        return f"{safe_title}.docx"

    def _generate_with_html2docx(self, article: Article, **options) -> bytes:
        """使用 html2docx 生成文档"""
        include_summary = options.get("include_summary", True)

        # 构建完整的 HTML 文档
        html_content = self._build_full_html(article, include_summary)

        # 预处理 HTML：下载图片并转换为 base64
        html_content = self._preprocess_images(html_content, str(article.url))

        # 使用 html2docx 转换
        buf = html2docx(html_content, title=article.title)
        value = buf.getvalue()
        if isinstance(value, str):
            return value.encode("utf-8")
        return bytes(value)

    def _build_full_html(self, article: Article, include_summary: bool = True) -> str:
        """构建完整的 HTML 文档"""
        # 构建元信息
        meta_items = []
        if article.account_name:
            meta_items.append(f"公众号: {article.account_name}")
        if article.author:
            meta_items.append(f"作者: {article.author}")
        if article.publish_time:
            meta_items.append(f"发布时间: {article.publish_time_str}")
        meta_items.append(f"字数: {article.word_count}")
        meta_html = " | ".join(meta_items)

        # 构建摘要部分
        summary_html = ""
        if include_summary and article.summary:
            key_points_html = ""
            if article.summary.key_points:
                points = "".join(f"<li>{p}</li>" for p in article.summary.key_points)
                key_points_html = f"<h3>📌 关键要点</h3><ul>{points}</ul>"

            tags_html = ""
            if article.summary.tags:
                tags = " ".join(
                    f"<span style='color: #07C160;'>#{t}</span>" for t in article.summary.tags
                )
                tags_html = f"<p><strong>标签:</strong> {tags}</p>"

            summary_html = f"""
            <div style="background-color: #f8f9fa; padding: 15px; margin: 20px 0; border-radius: 8px;">
                <h2>📝 文章摘要</h2>
                <p>{article.summary.content}</p>
                {key_points_html}
                {tags_html}
            </div>
            <hr/>
            """

        # 获取正文 HTML
        content_html = article.content_html or f"<p>{article.content_text}</p>"

        # 构建完整 HTML
        full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{article.title}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.8;
            color: #333;
        }}
        h1 {{
            text-align: center;
            font-size: 18pt;
            color: #1a1a1a;
        }}
        h2 {{
            font-size: 14pt;
            color: #07C160;
        }}
        h3 {{
            font-size: 12pt;
        }}
        .meta {{
            text-align: center;
            color: #666;
            font-size: 10pt;
        }}
        .source-link {{
            text-align: center;
            font-size: 9pt;
            color: #07C160;
        }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 10px auto;
        }}
        p {{
            margin-bottom: 1em;
        }}
        blockquote {{
            margin-left: 20px;
            padding-left: 15px;
            border-left: 3px solid #07C160;
            color: #666;
        }}
        .footer {{
            text-align: center;
            font-size: 8pt;
            color: #999;
            margin-top: 30px;
            border-top: 1px solid #eee;
            padding-top: 20px;
        }}
    </style>
</head>
<body>
    <h1>{article.title}</h1>
    <p class="meta">{meta_html}</p>
    <p class="source-link">原文链接: <a href="{article.url!s}">{article.url!s}</a></p>
    <hr/>
    {summary_html}
    <h2>📄 正文内容</h2>
    {content_html}
    <div class="footer">
        <p>文章ID: {article.id}</p>
        <p>抓取时间: {article.created_at.strftime("%Y-%m-%d %H:%M:%S")}</p>
        <p>由 WeChat Article Summarizer 生成</p>
    </div>
</body>
</html>"""
        return full_html

    def _preprocess_images(self, html_content: str, base_url: str) -> str:
        """预处理图片：下载并转换为 base64 内嵌"""
        soup = BeautifulSoup(html_content, "lxml")

        for img in soup.find_all("img"):
            raw_img_url = img.get("data-src") or img.get("src")
            if not raw_img_url:
                continue
            img_url = str(raw_img_url)

            # 跳过已经是 base64 的图片
            if img_url.startswith("data:"):
                continue

            # 跳过表情包
            if "emoji" in img_url.lower() or "emotion" in img_url.lower():
                img.decompose()
                continue

            # 处理相对 URL
            if not img_url.startswith(("http://", "https://")):
                img_url = urljoin(base_url, img_url)

            try:
                response = safe_fetch_sync(
                    img_url,
                    headers={
                        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                        "Referer": base_url,
                    },
                    timeout=30,
                )
                response.raise_for_status()

                content_type = response.headers.get("content-type", "image/png")
                if not content_type.startswith("image/"):
                    img.decompose()
                    continue
                if len(response.content) > MAX_REMOTE_IMAGE_BYTES:
                    raise ExporterError("图片体积超出安全限制")

                # 转换为 base64
                img_base64 = base64.b64encode(response.content).decode("ascii")
                img["src"] = f"data:{content_type};base64,{img_base64}"

                # 移除 data-src 属性
                if img.get("data-src"):
                    del img["data-src"]

                # 设置图片样式
                img["style"] = "max-width: 100%; height: auto; display: block; margin: 10px auto;"

                logger.debug(f"图片转换成功: {img_url[:50]}...")

            except Exception as e:
                logger.warning(f"图片下载失败 ({img_url[:50]}...): {e}")
                # 替换为占位符
                placeholder = soup.new_tag("p")
                placeholder.string = "[图片加载失败]"
                placeholder["style"] = "text-align: center; color: #ccc; font-size: 9pt;"
                img.replace_with(placeholder)

        return str(soup)

    def _generate_document_manual(self, article: Article, **options) -> DocxDocument:
        """手动构建 Word 文档（备用方案）"""
        include_summary = options.get("include_summary", True)

        doc = Document()

        # 设置Word 2013+兼容性模式，避免在Microsoft 365中显示为兼容模式
        self._set_compatibility_mode(doc)

        self._setup_styles(doc)

        # 标题
        title_para = doc.add_heading(article.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_items = []
        if article.account_name:
            meta_items.append(f"公众号: {article.account_name}")
        if article.author:
            meta_items.append(f"作者: {article.author}")
        if article.publish_time:
            meta_items.append(f"发布时间: {article.publish_time_str}")
        meta_items.append(f"字数: {article.word_count}")

        meta_run = meta_para.add_run(" | ".join(meta_items))
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        # 原文链接
        link_para = doc.add_paragraph()
        link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_run = link_para.add_run(f"原文链接: {article.url!s}")
        link_run.font.size = Pt(9)
        link_run.font.color.rgb = RGBColor(7, 193, 96)

        doc.add_paragraph("─" * 50)

        # 摘要
        if include_summary and article.summary:
            doc.add_heading("📝 文章摘要", level=1)
            summary_para = doc.add_paragraph(article.summary.content)
            summary_para.paragraph_format.first_line_indent = Pt(24)

            if article.summary.key_points:
                doc.add_paragraph()
                doc.add_heading("📌 关键要点", level=2)
                for point in article.summary.key_points:
                    doc.add_paragraph(point, style="List Bullet")

            doc.add_paragraph("─" * 50)

        # 正文
        doc.add_heading("📄 正文内容", level=1)
        self._add_content_with_images(doc, article)

        # 页脚
        doc.add_paragraph()
        doc.add_paragraph("─" * 50)

        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"文章ID: {article.id} | 抓取时间: {article.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(180, 180, 180)

        tool_para = doc.add_paragraph()
        tool_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tool_run = tool_para.add_run("由 WeChat Article Summarizer 生成")
        tool_run.font.size = Pt(8)
        tool_run.font.color.rgb = RGBColor(200, 200, 200)

        return doc

    def _add_content_with_images(self, doc: DocxDocument, article: Article) -> None:
        """添加正文内容（手动模式）"""
        content_html = article.content_html
        if not content_html:
            for para_text in article.content_text.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.first_line_indent = Pt(24)
            return

        soup = BeautifulSoup(content_html, "lxml")
        base_url = str(article.url)
        processed_images: set[str | int] = set()

        content_container = (
            soup.find(id="js_content")
            or soup.find(class_="rich_media_content")
            or soup.body
            or soup
        )
        self._process_children(doc, content_container, base_url, processed_images)

    def _process_children(
        self,
        doc: DocxDocument,
        container,
        base_url: str,
        processed_images: set[str | int],
    ) -> None:
        """递归处理子元素"""
        from bs4 import NavigableString

        for child in container.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text and len(text) > 1:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(8)
                continue

            tag_name = getattr(child, "name", None)
            if not tag_name or tag_name in ["script", "style", "meta", "link", "noscript"]:
                continue

            try:
                self._process_element(doc, child, base_url, processed_images)
            except Exception as e:
                logger.debug(f"处理元素失败: {e}")

    def _process_element(
        self,
        doc: DocxDocument,
        element,
        base_url: str,
        processed_images: set[str | int],
    ) -> None:
        """处理单个元素"""
        tag_name = element.name

        if tag_name == "img":
            img_id = element.get("data-src") or element.get("src") or id(element)
            if img_id not in processed_images:
                self._add_image(doc, element, base_url)
                processed_images.add(img_id)
            return

        if tag_name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            text = element.get_text(strip=True)
            if text:
                level = int(tag_name[1])
                doc.add_heading(text, level=min(level + 1, 9))
            return

        if tag_name == "p":
            # 先处理图片
            for img in element.find_all("img"):
                img_id = img.get("data-src") or img.get("src") or id(img)
                if img_id not in processed_images:
                    self._add_image(doc, img, base_url)
                    processed_images.add(img_id)
            # 再处理文本
            text = element.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.space_after = Pt(8)
            return

        if tag_name in ["section", "div", "article"]:
            self._process_children(doc, element, base_url, processed_images)
            return

        if tag_name in ["ul", "ol"]:
            for li in element.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                if text:
                    style = "List Bullet" if tag_name == "ul" else "List Number"
                    doc.add_paragraph(text, style=style)
            return

        if tag_name == "blockquote":
            text = element.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.left_indent = Pt(36)
                if para.runs:
                    para.runs[0].font.italic = True
                    para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
            return

        # 处理表格
        if tag_name == "table":
            self._add_table(doc, element)
            return

        if hasattr(element, "children"):
            self._process_children(doc, element, base_url, processed_images)

    def _add_image(self, doc: DocxDocument, img_element, base_url: str) -> None:
        """下载并添加图片"""
        raw_img_url = img_element.get("data-src") or img_element.get("src")
        if not raw_img_url:
            return
        img_url = str(raw_img_url)
        if img_url.startswith("data:"):
            return

        if "emoji" in img_url.lower() or "emotion" in img_url.lower():
            return

        if not img_url.startswith(("http://", "https://")):
            img_url = urljoin(base_url, img_url)

        try:
            response = safe_fetch_sync(
                img_url,
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
                    "Referer": base_url,
                },
                timeout=30,
            )
            response.raise_for_status()

            content_type = response.headers.get("content-type", "")
            if not content_type.startswith("image/"):
                return

            image_data = response.content
            if len(image_data) > MAX_REMOTE_IMAGE_BYTES:
                raise ExporterError("图片体积超出安全限制")
            image_stream = io.BytesIO(image_data)

            # 获取图片尺寸
            img_width_inches = self._get_image_width_inches(image_data)
            max_width_inches = 6.0

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()

            if img_width_inches and img_width_inches < max_width_inches:
                run.add_picture(image_stream)
            else:
                run.add_picture(image_stream, width=Inches(max_width_inches))

            doc.add_paragraph()

        except Exception as e:
            logger.warning(f"图片下载失败: {e}")
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("[图片加载失败]")
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(9)

    def _get_image_width_inches(self, image_data: bytes) -> float | None:
        """获取图片宽度（英寸）"""
        try:
            from PIL import Image

            image_stream = io.BytesIO(image_data)
            with Image.open(image_stream) as img:
                width_px = img.width
                dpi = img.info.get("dpi", (96, 96))
                dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
                if not dpi_x:
                    return None
                return float(width_px) / float(dpi_x)
        except Exception:
            return None

    def _add_table(self, doc: DocxDocument, table_element) -> None:
        """添加表格到Word文档"""
        try:
            # 提取所有行
            rows = table_element.find_all("tr")
            if not rows:
                return

            # 计算最大列数
            max_cols = 0
            for row in rows:
                cells = row.find_all(["td", "th"])
                col_count = sum(int(cell.get("colspan", 1)) for cell in cells)
                max_cols = max(max_cols, col_count)

            if max_cols == 0:
                return

            # 创建Word表格
            word_table = doc.add_table(rows=len(rows), cols=max_cols)
            word_table.style = "Table Grid"

            # 填充表格内容
            for row_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                col_idx = 0

                for cell in cells:
                    if col_idx >= max_cols:
                        break

                    # 获取单元格文本
                    cell_text = cell.get_text(strip=True)

                    # 设置单元格内容
                    word_cell = word_table.cell(row_idx, col_idx)
                    word_cell.text = cell_text

                    # 如果是表头，设置加粗
                    if cell.name == "th" or row_idx == 0:
                        for paragraph in word_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # 设置字体
                    for paragraph in word_cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = "微软雅黑"

                    # 处理合并单元格
                    colspan = int(cell.get("colspan", 1))
                    if colspan > 1 and col_idx + colspan <= max_cols:
                        for merge_idx in range(1, colspan):
                            word_table.cell(row_idx, col_idx).merge(
                                word_table.cell(row_idx, col_idx + merge_idx)
                            )

                    col_idx += colspan

            # 表格后添加空行
            doc.add_paragraph()
            logger.debug(f"添加表格: {len(rows)}行 x {max_cols}列")

        except Exception as e:
            logger.warning(f"表格处理失败: {e}")
            # 回退到纯文本输出
            text = table_element.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.paragraph_format.space_after = Pt(8)

    def _set_compatibility_mode(self, doc: DocxDocument) -> None:
        """设置文档兼容性模式为Word 2013+

        通过在settings.xml中设置compatibilityMode=15，
        确保文档在Microsoft 365中不会显示为兼容性模式。

        参考: https://learn.microsoft.com/en-us/openspecs/office_standards/ms-docx/
        compatibilityMode val="15" = Word 2013+
        """
        try:
            # 获取settings部分
            settings = doc.settings
            settings_element = settings._settings

            # 查找或创建w:compat元素
            compat = settings_element.find(qn("w:compat"))
            if compat is None:
                # 创建compat元素
                compat = etree.SubElement(settings_element, qn("w:compat"))

            # 查找现有的compatibilityMode设置
            compat_settings = compat.findall(qn("w:compatSetting"))
            mode_setting = None

            for cs in compat_settings:
                name = cs.get(qn("w:name"))
                uri = cs.get(qn("w:uri"))
                if (
                    name == "compatibilityMode"
                    and uri == "http://schemas.microsoft.com/office/word"
                ):
                    mode_setting = cs
                    break

            if mode_setting is not None:
                # 更新现有设置
                mode_setting.set(qn("w:val"), WORD_COMPATIBILITY_MODE_2013)
            else:
                # 创建新的compatSetting元素
                compat_setting = etree.SubElement(compat, qn("w:compatSetting"))
                compat_setting.set(qn("w:name"), "compatibilityMode")
                compat_setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
                compat_setting.set(qn("w:val"), WORD_COMPATIBILITY_MODE_2013)

                # 添加其他Word 2013+推荐的兼容性设置
                self._add_compat_setting(compat, "overrideTableStyleFontSizeAndJustification", "1")
                self._add_compat_setting(compat, "enableOpenTypeFeatures", "1")
                self._add_compat_setting(compat, "doNotFlipMirrorIndents", "1")
                self._add_compat_setting(compat, "differentiateMultirowTableHeaders", "1")

            logger.debug(f"已设置Word兼容性模式为 Word 2013+ (val={WORD_COMPATIBILITY_MODE_2013})")

        except Exception as e:
            logger.warning(f"设置兼容性模式失败: {e}")

    def _add_compat_setting(self, compat_element, name: str, value: str) -> None:
        """添加兼容性设置项"""
        setting = etree.SubElement(compat_element, qn("w:compatSetting"))
        setting.set(qn("w:name"), name)
        setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
        setting.set(qn("w:val"), value)

    def _setup_styles(self, doc: DocxDocument) -> None:
        """设置文档样式"""
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)

        # 设置中文字体
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        for i in range(1, 4):
            heading_style = doc.styles[f"Heading {i}"]
            heading_style.font.name = "微软雅黑"
            heading_style.font.bold = True
            heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
