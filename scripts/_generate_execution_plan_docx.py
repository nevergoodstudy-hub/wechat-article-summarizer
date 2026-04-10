from __future__ import annotations

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

OUT = r"D:\\Newidea-warp\\项目改进任务可执行方案（2026-开发组执行版）.docx"


def set_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)


def add_cover(doc: Document) -> None:
    doc.add_heading("微信文章总结器项目改进任务可执行方案（2026）", 0)
    p = doc.add_paragraph()
    p.add_run("角色视角：大型软件项目开发组（架构组/安全组/测试组/平台组/产品组）\n")
    p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
    p.add_run("输入依据：\n")
    p.add_run("1) 《下一代深度改进方针（2026）》\n")
    p.add_run("2) 《项目下一代深度改进方针（2026-深度刨析版）》\n")
    p.add_run("3) 《下一代改进计划-任务分解WBS与里程碑甘特（2026）》\n")
    p.add_run("4) 本地代码与目录审阅 + 外部最佳实践（Python TaskGroup/OWASP SSRF/pytest fixtures/structlog）")


def add_summary(doc: Document) -> None:
    doc.add_heading("一、执行总览（先做什么、谁负责、如何验收）", level=1)
    bullets = [
        "总目标：在 12 个月内完成‘安全收敛—架构重整—质量固化—平台扩展’四阶段交付闭环。",
        "执行原则：P0/P1 先行、风险前置、每项任务必须绑定 Owner / 截止时间 / 验收证据。",
        "治理策略：将‘方针文件’转化为可执行工作包（WP），每个工作包最小可交付周期为 2 周。",
        "验收方式：代码PR + 自动化检查（lint/type/test/security）+ 运行指标 + 文档归档 四证齐全。",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def add_sources(doc: Document) -> None:
    doc.add_heading("二、联网与技能依据（开发组采用）", level=1)
    doc.add_paragraph("已结合项目安全/审计类技能框架，外部依据聚焦以下可落地规范：")
    refs = [
        "Python 3.14 asyncio 文档：TaskGroup 为结构化并发推荐路径。",
        "OWASP SSRF Cheat Sheet：URL/域名/IP校验与禁止自动重定向是关键控制点。",
        "pytest fixtures 文档：通过显式 fixture 与作用域治理提升测试可重复性与独立性。",
        "structlog 文档：结构化日志 + 上下文字段治理，支撑审计与追踪。",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")


def add_risk_to_workpackage(doc: Document) -> None:
    doc.add_heading("三、问题到任务映射（Issue → Work Package）", level=1)
    table = doc.add_table(rows=1, cols=6)
    h = table.rows[0].cells
    h[0].text = "问题级别"
    h[1].text = "核心问题"
    h[2].text = "执行工作包"
    h[3].text = "Owner"
    h[4].text = "周期"
    h[5].text = "验收标准"

    rows = [
        ("P0", "MCP高风险操作面", "WP-SEC-01 输入校验统一层 + 高危操作二次确认", "安全组", "W1-W4", "注入/越权回归用例100%通过"),
        ("P0", "GUI主文件体量过大", "WP-ARC-01 GUI分层拆分（Frame/ViewModel）", "架构组", "W1-W8", "app.py 体量下降>60%，核心流程回归通过"),
        ("P1", "并发模型不统一", "WP-REL-01 gather→TaskGroup迁移", "平台组", "W3-W8", "并发异常聚合可观测，无孤儿任务"),
        ("P1", "测试隔离不足", "WP-QA-01 测试容器与fixture重构", "测试组", "W1-W6", "随机顺序执行通过率≥95%"),
        ("P1", "日志敏感字段泄露风险", "WP-SEC-02 审计日志脱敏标准化", "安全组", "W2-W5", "敏感字段扫描0泄露"),
        ("P2", "类型与质量债务", "WP-ENG-01 mypy/ruff债务清理", "平台组", "W5-W12", "新增代码类型覆盖达标，ruff关键规则清零"),
        ("P2", "CI门禁不完整", "WP-ENG-02 质量流水线闭环", "平台组", "W4-W10", "PR门禁含lint/type/test/security/build"),
        ("P3", "文档与指标碎片化", "WP-OPS-01 质量看板与版本归档机制", "产品组", "W9-W14", "每版本自动产出质量报告"),
    ]

    for r in rows:
        c = table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def add_wbs_detail(doc: Document) -> None:
    doc.add_heading("四、可执行WBS（细化到行动级）", level=1)
    table = doc.add_table(rows=1, cols=7)
    h = table.rows[0].cells
    h[0].text = "WBS"
    h[1].text = "任务"
    h[2].text = "输入"
    h[3].text = "输出"
    h[4].text = "负责人"
    h[5].text = "前置条件"
    h[6].text = "DoD（完成定义）"

    tasks = [
        ("1.1", "建立MCP参数Schema与校验网关", "现有server.py/input_validator.py", "统一validator模块+拒绝策略", "安全组", "接口清单冻结", "所有工具参数进入统一校验路径"),
        ("1.2", "高危工具人机确认(HITL)", "权限分级矩阵", "确认流程与审计记录", "安全组", "管理员角色定义", "高危操作未确认不可执行"),
        ("2.1", "GUI职责解耦切分", "app.py职责清单", "Sidebar/List/Summary/Export组件", "架构组", "UI契约定义", "核心页面组件化并可独立测试"),
        ("2.2", "ViewModel接口固化", "现有事件处理逻辑", "ViewModel接口文档", "架构组", "UseCase入参统一", "UI逻辑与业务逻辑分离"),
        ("3.1", "批处理并发迁移TaskGroup", "batch流程与异常路径", "TaskGroup实现与except*处理", "平台组", "Python版本基线", "并发任务失败可控取消"),
        ("3.2", "统一超时/重试/退避策略", "现有http/llm调用点", "策略中间层", "平台组", "调用链梳理完成", "重试行为可配置且可观测"),
        ("4.1", "测试容器与mock边界重构", "tests/conftest.py", "测试专用Container", "测试组", "端口接口清单", "单测可独立运行"),
        ("4.2", "随机顺序与并行稳定性治理", "pytest配置", "随机序回归与隔离修复", "测试组", "基础失败用例收敛", "重跑稳定性达标"),
        ("5.1", "结构化日志敏感字段脱敏", "日志字段盘点", "脱敏处理器", "安全组", "日志规范确定", "密钥/token/密码不落盘"),
        ("6.1", "CI质量门禁编排", ".github/workflows", "统一pipeline", "平台组", "执行时长预算", "门禁失败可阻断合并"),
    ]

    for t in tasks:
        c = table.add_row().cells
        for i, v in enumerate(t):
            c[i].text = v


def add_milestone_gantt(doc: Document) -> None:
    doc.add_heading("五、里程碑与直观甘特（季度视图）", level=1)
    doc.add_paragraph("图例：█ 执行期  · 空档")

    gantt = [
        ("M1 安全基线收敛（W1-W4）", "████ ···· ···· ····"),
        ("M2 GUI架构重整（W1-W8）", "████ ████ ···· ····"),
        ("M3 并发与可靠性升级（W3-W8）", "··██ ████ ···· ····"),
        ("M4 测试隔离专项（W1-W6）", "████ ██·· ···· ····"),
        ("M5 质量门禁闭环（W4-W10）", "···█ ████ ██·· ····"),
        ("M6 平台化与看板（W9-W14）", "···· ···· ████ ██··"),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "里程碑"
    table.rows[0].cells[1].text = "W1-W16"
    for name, bar in gantt:
        c = table.add_row().cells
        c[0].text = name
        c[1].text = bar


def add_raci(doc: Document) -> None:
    doc.add_heading("六、RACI责任矩阵（团队协同）", level=1)
    table = doc.add_table(rows=1, cols=6)
    h = table.rows[0].cells
    h[0].text = "工作包"
    h[1].text = "架构组"
    h[2].text = "安全组"
    h[3].text = "测试组"
    h[4].text = "平台组"
    h[5].text = "产品组"

    raci = [
        ("WP-SEC-01", "C", "A/R", "C", "C", "I"),
        ("WP-ARC-01", "A/R", "C", "C", "C", "I"),
        ("WP-REL-01", "C", "C", "C", "A/R", "I"),
        ("WP-QA-01", "C", "C", "A/R", "C", "I"),
        ("WP-ENG-02", "I", "C", "C", "A/R", "C"),
        ("WP-OPS-01", "I", "I", "C", "C", "A/R"),
    ]

    for row in raci:
        c = table.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v


def add_acceptance(doc: Document) -> None:
    doc.add_heading("七、每两周冲刺执行模板（可直接落地）", level=1)
    steps = [
        "Sprint Planning：按风险优先级选择2~4个WBS任务，冻结范围。",
        "Daily Execution：每天同步阻塞项，关键任务挂风险标签（security/reliability）。",
        "Mid-Sprint Check：检查自动化门禁通过率与缺陷趋势，必要时动态降级范围。",
        "Sprint Review：提交‘代码+测试+报告’三件套并做演示。",
        "Sprint Retro：复盘缺陷根因，更新下一迭代纠偏动作。",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("八、验收与发布闸门（Release Gate）", level=1)
    gate_table = doc.add_table(rows=1, cols=4)
    gate_table.rows[0].cells[0].text = "闸门"
    gate_table.rows[0].cells[1].text = "必备检查"
    gate_table.rows[0].cells[2].text = "通过阈值"
    gate_table.rows[0].cells[3].text = "证据"
    gates = [
        ("G1 安全闸门", "注入/越权/SSRF回归", "高危0遗留", "安全回归报告"),
        ("G2 质量闸门", "ruff+mypy+pytest", "主分支通过率≥95%", "CI记录"),
        ("G3 架构闸门", "边界扫描+模块规模", "跨层违规=0", "架构检查报告"),
        ("G4 运营闸门", "日志脱敏+追踪", "敏感信息泄露=0", "审计日志抽检"),
    ]
    for g in gates:
        c = gate_table.add_row().cells
        for i, v in enumerate(g):
            c[i].text = v


def add_closing(doc: Document) -> None:
    doc.add_heading("九、结论", level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        "该可执行方案已将方针文件和深度剖析文件转化为可排期、可分工、可验收的执行单元。"
        "建议从 WP-SEC-01、WP-ARC-01、WP-QA-01 同步启动，形成‘安全+架构+测试’三线并行推进，"
        "以 4 周达成首个可验证里程碑。"
    )
    run.font.color.rgb = RGBColor(32, 32, 32)


def main() -> None:
    doc = Document()
    set_font(doc)
    add_cover(doc)
    add_summary(doc)
    add_sources(doc)
    add_risk_to_workpackage(doc)
    add_wbs_detail(doc)
    add_milestone_gantt(doc)
    add_raci(doc)
    add_acceptance(doc)
    add_closing(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
