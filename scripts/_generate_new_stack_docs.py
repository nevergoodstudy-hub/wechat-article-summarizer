from __future__ import annotations

from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = r"D:\\Newidea-warp"
DOC1 = ROOT + r"\新技术栈深度改进实施方案（2026）.docx"
DOC2 = ROOT + r"\新技术栈迁移WBS与里程碑甘特（可执行版）.docx"


def _set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)


def build_doc1() -> None:
    doc = Document()
    _set_style(doc)

    doc.add_heading("新技术栈深度改进实施方案（2026）", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("视角：大型软件项目开发团队（架构组/平台组/安全组/客户端组/测试组）")

    doc.add_heading("一、目标与约束（功能完全等价）", level=1)
    for x in [
        "必须保持目标与功能一致：抓取、摘要、RAG/GraphRAG、导出、GUI/CLI/MCP、缓存、审计与CI。",
        "替换现有主栈，构建可长期维护、可观测、可审计、可扩展的新平台。",
        "迁移期间采用双轨验证，避免业务中断。",
    ]:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("二、新技术栈选型（2026）", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    h = tbl.rows[0].cells
    h[0].text = "能力域"
    h[1].text = "现栈"
    h[2].text = "新栈"
    h[3].text = "选择理由"
    rows = [
        ("桌面GUI", "CustomTkinter", "Tauri 2 + React + Vite + TypeScript", "跨平台、体积小、安全边界更强"),
        ("核心后端", "Python单体", "Rust (Axum) + OpenAPI", "性能与并发可靠性更高，接口契约清晰"),
        ("抓取", "httpx/playwright", "reqwest + 渲染服务", "链路隔离，便于安全与限流治理"),
        ("RAG", "Chroma + sentence-transformers", "Qdrant + FastEmbed/BGE", "向量检索能力更强，工程化成熟"),
        ("GraphRAG", "networkx/igraph", "petgraph + 图算法服务", "更适合高性能服务化"),
        ("MCP", "Python MCP服务", "Rust MCP独立进程", "权限与审计策略可内建、可硬化"),
        ("可观测", "日志为主", "OpenTelemetry + tracing + Prometheus", "全链路追踪、容量规划、故障定位"),
    ]
    for r in rows:
        c = tbl.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading("三、目标架构", level=1)
    for x in [
        "core-domain：纯领域模型与规则。",
        "core-application：用例编排与端口定义。",
        "adapters：抓取器、LLM、导出器、向量库、存储。",
        "platform-gateway：统一API（GUI/CLI/MCP共用）。",
        "desktop-shell：Tauri前端壳层。",
    ]:
        doc.add_paragraph(x, style="List Number")

    doc.add_heading("四、迁移实施路径（12个月）", level=1)
    roadmap = doc.add_table(rows=1, cols=3)
    roadmap.rows[0].cells[0].text = "阶段"
    roadmap.rows[0].cells[1].text = "周期"
    roadmap.rows[0].cells[2].text = "关键交付"
    for s in [
        ("Phase 1 基础骨架", "M1-M2", "Rust workspace、OpenAPI、统一DTO、CLI骨架"),
        ("Phase 2 核心能力迁移", "M3-M5", "抓取/摘要/导出链路等价实现"),
        ("Phase 3 智能能力", "M6-M8", "RAG/GraphRAG 与 MCP 独立服务"),
        ("Phase 4 客户端替换", "M9-M10", "Tauri GUI 完成旧功能映射"),
        ("Phase 5 双轨验收切换", "M11-M12", "A/B验证、压测、安全验收、正式切换"),
    ]:
        c = roadmap.add_row().cells
        c[0].text, c[1].text, c[2].text = s

    doc.add_heading("五、验收标准", level=1)
    for x in [
        "功能一致性：关键业务回归通过率 100%。",
        "质量门禁：lint/type/test/security/build 全通过。",
        "安全门禁：MCP注入/越权/审计泄露高危为0。",
        "性能门禁：批量吞吐和稳定性达到目标基线。",
    ]:
        doc.add_paragraph(x, style="List Bullet")

    doc.save(DOC1)


def build_doc2() -> None:
    doc = Document()
    _set_style(doc)

    doc.add_heading("新技术栈迁移WBS与里程碑甘特（可执行版）", 0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("一、WBS任务分解", level=1)
    wbs = doc.add_table(rows=1, cols=7)
    h = wbs.rows[0].cells
    h[0].text = "WBS"
    h[1].text = "任务"
    h[2].text = "Owner"
    h[3].text = "输入"
    h[4].text = "输出"
    h[5].text = "前置"
    h[6].text = "DoD"

    items = [
        ("1.1", "Rust核心服务框架搭建", "平台组", "现有接口清单", "Axum服务骨架", "架构评审", "服务可启动+健康检查"),
        ("1.2", "OpenAPI契约冻结", "架构组", "现有DTO", "API契约文档", "1.1", "GUI/CLI/MCP共用契约"),
        ("2.1", "抓取链路迁移", "平台组", "旧抓取逻辑", "新抓取适配器", "1.2", "主站点抓取通过"),
        ("2.2", "摘要与多Provider迁移", "算法组", "旧摘要器", "统一LLM网关", "1.2", "方法等价可用"),
        ("2.3", "导出链路迁移", "平台组", "旧导出器", "HTML/MD/DOCX导出", "1.2", "导出回归通过"),
        ("3.1", "Qdrant RAG集成", "算法组", "向量方案", "检索增强摘要", "2.2", "效果评估达标"),
        ("3.2", "GraphRAG服务实现", "算法组", "图分析需求", "图谱摘要服务", "3.1", "复杂问题回答可用"),
        ("4.1", "MCP独立服务重建", "安全组", "旧MCP工具集", "新MCP服务", "1.2", "工具调用回归通过"),
        ("4.2", "MCP安全硬化", "安全组", "安全策略", "校验/审计/确认/白名单", "4.1", "高危用例全拦截"),
        ("5.1", "Tauri GUI替换", "客户端组", "旧GUI流程", "新GUI壳层", "1.2", "关键流程可用"),
        ("5.2", "双轨对比验收", "测试组", "新旧系统", "对比报告", "2.x/3.x/4.x/5.1", "一致性100%"),
        ("6.1", "发布切换与回滚预案", "平台组", "验收报告", "上线方案", "5.2", "灰度切换成功"),
    ]
    for it in items:
        c = wbs.add_row().cells
        for i, v in enumerate(it):
            c[i].text = v

    doc.add_heading("二、里程碑甘特（文本可视化）", level=1)
    doc.add_paragraph("图例：█ 执行期  · 空档")
    g = doc.add_table(rows=1, cols=2)
    g.rows[0].cells[0].text = "里程碑"
    g.rows[0].cells[1].text = "M1-M12"
    for row in [
        ("M1 平台骨架与契约", "███·········"),
        ("M2 抓取/摘要/导出迁移", "··███·······"),
        ("M3 RAG/GraphRAG", "····███·····"),
        ("M4 MCP重建与硬化", "·····███····"),
        ("M5 Tauri GUI替换", "·······███··"),
        ("M6 双轨验收与切换", "·········███"),
    ]:
        c = g.add_row().cells
        c[0].text, c[1].text = row

    doc.add_heading("三、发布闸门", level=1)
    gates = doc.add_table(rows=1, cols=4)
    gates.rows[0].cells[0].text = "闸门"
    gates.rows[0].cells[1].text = "检查项"
    gates.rows[0].cells[2].text = "阈值"
    gates.rows[0].cells[3].text = "证据"
    for x in [
        ("G1 功能", "关键用例回归", "100%通过", "回归报告"),
        ("G2 安全", "注入/越权/审计泄露", "高危0", "安全测试报告"),
        ("G3 性能", "吞吐/延迟/稳定性", "达成基线", "压测报告"),
        ("G4 工程", "CI质量门禁", "主干通过率>=95%", "流水线记录"),
    ]:
        c = gates.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = x

    doc.save(DOC2)


if __name__ == "__main__":
    build_doc1()
    build_doc2()
    print(DOC1)
    print(DOC2)
