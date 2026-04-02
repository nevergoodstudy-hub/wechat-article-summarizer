from __future__ import annotations

from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = r"D:\\Newidea-warp"
DOC_DIR = ROOT + r"\docs\new_stack_program"


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)


def save_doc(filename: str, builder) -> str:
    path = DOC_DIR + "\\" + filename
    doc = Document()
    set_style(doc)
    builder(doc)
    doc.save(path)
    return path


def build_arch_decisions(doc: Document) -> None:
    doc.add_heading("架构决策记录 ADR（新技术栈）", 0)
    doc.add_paragraph(f"版本：v1.0\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    decisions = [
        ("ADR-001", "采用 Rust + TypeScript 双栈", "提升性能与工程化边界，统一GUI/API/MCP调用契约"),
        ("ADR-002", "桌面端选型 Tauri 2 + React", "替代传统桌面UI技术，降低包体并增强安全隔离"),
        ("ADR-003", "MCP 服务独立进程化", "最小权限运行、可审计、可回滚"),
        ("ADR-004", "统一 OpenAPI 契约驱动", "CLI/GUI/MCP 共享业务接口，减少重复逻辑"),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "编号"
    t.rows[0].cells[1].text = "决策"
    t.rows[0].cells[2].text = "理由"
    t.rows[0].cells[3].text = "状态"
    for d in decisions:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = d
        c[3].text = "Approved"


def build_requirements_trace(doc: Document) -> None:
    doc.add_heading("需求追踪矩阵 RTM（旧功能 → 新栈实现）", 0)
    doc.add_paragraph("确保‘目的与功能完全等价’。")
    t = doc.add_table(rows=1, cols=5)
    t.rows[0].cells[0].text = "需求ID"
    t.rows[0].cells[1].text = "旧系统功能"
    t.rows[0].cells[2].text = "新系统实现"
    t.rows[0].cells[3].text = "验证方式"
    t.rows[0].cells[4].text = "状态"
    rows = [
        ("REQ-01", "文章抓取", "Scraper Adapter + 渲染服务", "站点回归用例", "Planned"),
        ("REQ-02", "摘要生成", "LLM Gateway", "同输入输出对比", "Planned"),
        ("REQ-03", "RAG/GraphRAG", "Qdrant + graph service", "评测集对比", "Planned"),
        ("REQ-04", "导出", "Exporter Service", "格式回归", "Planned"),
        ("REQ-05", "MCP", "Rust MCP Server", "协议兼容测试", "Planned"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def build_api_spec(doc: Document) -> None:
    doc.add_heading("API 契约与版本策略（OpenAPI First）", 0)
    for p in [
        "统一网关：/v1/fetch, /v1/summarize, /v1/export, /v1/rag, /v1/graphrag",
        "版本策略：URL版本 + 向后兼容窗口（至少2个小版本）",
        "错误模型：统一错误码与可机器解析字段（code/message/details/correlation_id）",
        "发布规则：API变更必须附带契约测试与迁移说明",
    ]:
        doc.add_paragraph(p, style="List Bullet")


def build_security_baseline(doc: Document) -> None:
    doc.add_heading("安全基线与威胁模型（STRIDE + 零信任）", 0)
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "风险域"
    t.rows[0].cells[1].text = "控制要求"
    t.rows[0].cells[2].text = "验证"
    t.rows[0].cells[3].text = "门禁"
    rows = [
        ("SSRF", "单次解析绑定+重定向逐跳校验+host allowlist", "攻击样例回归", "High=0"),
        ("MCP 注入/越权", "参数schema+权限分层+HITL确认", "渗透脚本", "High=0"),
        ("日志泄露", "敏感字段脱敏与长度截断", "日志扫描", "Leak=0"),
        ("供应链", "依赖漏洞扫描与SBOM", "CI审计", "Critical=0"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def build_test_strategy(doc: Document) -> None:
    doc.add_heading("测试策略与验收矩阵", 0)
    for p in [
        "测试金字塔：单元 > 契约 > 集成 > E2E > 安全 > 性能",
        "门禁流水线：ruff/mypy(or equivalent)/unit/integration/security/build",
        "双轨验收：旧栈与新栈关键路径A/B比对",
        "发布闸门：功能一致性100%，主干通过率≥95%",
    ]:
        doc.add_paragraph(p, style="List Number")


def build_sre_slo(doc: Document) -> None:
    doc.add_heading("SRE运行手册与SLO/SLI", 0)
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "服务"
    t.rows[0].cells[1].text = "SLI"
    t.rows[0].cells[2].text = "SLO"
    t.rows[0].cells[3].text = "告警阈值"
    rows = [
        ("API网关", "请求成功率", "99.5%", "5分钟<99%"),
        ("摘要服务", "P95响应时间", "<3s(短文)", "P95>5s"),
        ("MCP服务", "工具调用成功率", "99.9%", "15分钟<99.5%"),
        ("导出服务", "任务完成率", "99.0%", "30分钟<98%"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def build_release_rollback(doc: Document) -> None:
    doc.add_heading("发布计划与回滚预案", 0)
    for p in [
        "发布策略：蓝绿/金丝雀，按用户组分批切换",
        "回滚触发：高危安全事件、核心链路错误率超阈值、数据一致性失败",
        "回滚目标：15分钟内恢复旧栈主业务可用",
        "演练制度：每季度至少一次全链路回滚演练",
    ]:
        doc.add_paragraph(p, style="List Bullet")


def build_risk_register(doc: Document) -> None:
    doc.add_heading("风险登记册与应对计划", 0)
    t = doc.add_table(rows=1, cols=5)
    t.rows[0].cells[0].text = "风险"
    t.rows[0].cells[1].text = "概率"
    t.rows[0].cells[2].text = "影响"
    t.rows[0].cells[3].text = "应对"
    t.rows[0].cells[4].text = "Owner"
    for r in [
        ("Rust学习曲线", "中", "高", "模板仓+评审机制+Pairing", "架构组"),
        ("功能等价回归不足", "中", "高", "RTM驱动 + 自动回归矩阵", "测试组"),
        ("MCP安全策略遗漏", "低", "高", "红队脚本+门禁阻断", "安全组"),
        ("迁移周期拉长", "中", "中", "里程碑拆小+周度风险审查", "PMO"),
    ]:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def build_governance(doc: Document) -> None:
    doc.add_heading("治理模型（RACI + 变更流程）", 0)
    doc.add_paragraph("RACI：架构组(A)、平台组(R)、安全组(R)、测试组(R)、产品组(C)、运维组(C)")
    doc.add_paragraph("变更流程：RFC -> 架构评审 -> 安全评审 -> 实施 -> 验收 -> 发布")


def build_training(doc: Document) -> None:
    doc.add_heading("团队培训与能力建设计划", 0)
    rows = [
        "Rust工程规范与异步编程",
        "Tauri桌面安全与前端工程化",
        "MCP协议实现与安全实践",
        "SRE监控与事件响应演练",
    ]
    for r in rows:
        doc.add_paragraph(r, style="List Number")


def build_procurement(doc: Document) -> None:
    doc.add_heading("资源与预算规划", 0)
    doc.add_paragraph("建议预算维度：人力（架构/平台/安全/测试）、CI算力、监控存储、代码签名证书、第三方API配额。")


def build_compliance(doc: Document) -> None:
    doc.add_heading("合规与数据治理", 0)
    for p in [
        "数据分类分级与最小化收集原则",
        "敏感信息生命周期管理（采集/存储/传输/销毁）",
        "审计留痕与可追溯性要求",
        "第三方服务合规条款审查",
    ]:
        doc.add_paragraph(p, style="List Bullet")


def build_docs_catalog(doc: Document) -> None:
    doc.add_heading("文档总目录与维护责任", 0)
    t = doc.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "文档"
    t.rows[0].cells[1].text = "责任团队"
    t.rows[0].cells[2].text = "更新频率"
    rows = [
        ("ADR", "架构组", "按重大变更"),
        ("RTM", "测试组", "每迭代"),
        ("安全基线", "安全组", "每月"),
        ("SRE手册", "运维组", "每月"),
        ("发布回滚预案", "平台组", "每次发布前"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def main() -> None:
    builders = [
        ("01-ADR-架构决策记录（新栈）.docx", build_arch_decisions),
        ("02-RTM-需求追踪矩阵（旧功能到新栈映射）.docx", build_requirements_trace),
        ("03-API契约与版本策略（OpenAPI）.docx", build_api_spec),
        ("04-安全基线与威胁模型（STRIDE）.docx", build_security_baseline),
        ("05-测试策略与验收矩阵.docx", build_test_strategy),
        ("06-SRE运行手册与SLO-SLI.docx", build_sre_slo),
        ("07-发布计划与回滚预案.docx", build_release_rollback),
        ("08-风险登记册与应对计划.docx", build_risk_register),
        ("09-治理模型与变更流程.docx", build_governance),
        ("10-团队培训与能力建设计划.docx", build_training),
        ("11-资源与预算规划.docx", build_procurement),
        ("12-合规与数据治理要求.docx", build_compliance),
        ("13-文档总目录与维护责任矩阵.docx", build_docs_catalog),
    ]

    paths: list[str] = []
    for fname, builder in builders:
        paths.append(save_doc(fname, builder))

    for p in paths:
        print(p)


if __name__ == "__main__":
    main()
