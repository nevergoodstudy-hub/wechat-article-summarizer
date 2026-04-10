from __future__ import annotations

from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = r"D:\\Newidea-warp"
OUT_DIR = ROOT + r"\docs\project_focused"


def style(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    s.font.size = Pt(10.5)


def mk(name: str, fn) -> str:
    p = OUT_DIR + "\\" + name
    d = Document()
    style(d)
    fn(d)
    d.save(p)
    return p


def d1_arch_maint(d: Document) -> None:
    d.add_heading("项目开发维护总指南（wechat-summarizer）", 0)
    d.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    d.add_heading("1. 代码目录与职责", 1)
    t = d.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "目录"
    t.rows[0].cells[1].text = "职责"
    t.rows[0].cells[2].text = "维护要点"
    rows = [
        ("src/wechat_summarizer/domain", "实体/值对象/领域服务", "禁止依赖 infrastructure/presentation"),
        ("src/wechat_summarizer/application", "用例与端口", "保持端口稳定，避免与具体适配器耦合"),
        ("src/wechat_summarizer/infrastructure", "适配器实现", "外部依赖异常统一封装、可降级"),
        ("src/wechat_summarizer/presentation", "CLI/GUI", "仅做交互编排，业务逻辑下沉用例"),
        ("src/wechat_summarizer/mcp", "MCP工具与安全", "输入校验、权限分级、审计脱敏必须开启"),
        ("tests", "单元/集成/回归", "新增功能必须同步测试与回归样例"),
    ]
    for r in rows:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = r

    d.add_heading("2. 维护红线", 1)
    for x in [
        "新增代码不得绕过 MCPInputValidator 与 SSRF 防护。",
        "禁止在 domain 层引入基础设施依赖。",
        "提交前必须通过 lint + test（至少受影响模块）。",
        "涉及导出/安全/抓取改动必须补充回归用例。",
    ]:
        d.add_paragraph(x, style="List Bullet")


def d2_module_tasks(d: Document) -> None:
    d.add_heading("模块级改进任务清单（面向当前仓库）", 0)
    d.add_paragraph("用途：开发者按模块领取任务，避免泛化改造。")
    t = d.add_table(rows=1, cols=5)
    t.rows[0].cells[0].text = "模块"
    t.rows[0].cells[1].text = "任务"
    t.rows[0].cells[2].text = "代码位置"
    t.rows[0].cells[3].text = "验收"
    t.rows[0].cells[4].text = "优先级"
    rows = [
        ("MCP", "补齐全部工具入参校验", "mcp/server.py", "恶意参数回归通过", "P0"),
        ("MCP", "日志敏感字段二次脱敏审计", "mcp/security.py", "抽检无密钥泄露", "P0"),
        ("抓取", "统一超时/重试/退避策略", "infrastructure/adapters/scrapers", "网络异常恢复可观测", "P1"),
        ("并发", "批处理链路 TaskGroup 化", "application/use_cases", "异常聚合正确", "P1"),
        ("GUI", "继续拆分 app.py 与事件总线收敛", "presentation/gui", "单文件体量下降、功能回归通过", "P1"),
        ("测试", "修复不可独立运行测试", "tests", "随机顺序通过率>=95%", "P1"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def d3_runbook(d: Document) -> None:
    d.add_heading("本项目开发运行手册（本地）", 0)
    d.add_heading("1. 最小启动路径", 1)
    for x in [
        "安装：pip install -e .",
        "CLI：python -m wechat_summarizer fetch <url>",
        "GUI：python -m wechat_summarizer",
        "MCP：python -m wechat_summarizer.mcp",
    ]:
        d.add_paragraph(x, style="List Number")
    d.add_heading("2. 故障定位优先顺序", 1)
    for x in [
        "先查输入校验与配置（API Key/URL/导出路径）。",
        "再查容器装配与适配器可用性。",
        "最后查外部依赖（网络、第三方 API、浏览器运行时）。",
    ]:
        d.add_paragraph(x, style="List Bullet")


def d4_security(d: Document) -> None:
    d.add_heading("本项目安全加固基线（开发必读）", 0)
    t = d.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "控制项"
    t.rows[0].cells[1].text = "当前实现位置"
    t.rows[0].cells[2].text = "开发要求"
    t.rows[0].cells[3].text = "验证方法"
    rows = [
        ("URL+SSRF 防护", "shared/utils/ssrf_protection.py", "新增网络调用必须复用安全客户端", "恶意URL回归"),
        ("MCP 输入校验", "mcp/input_validator.py", "每个工具入口都校验", "注入样例测试"),
        ("审计日志脱敏", "mcp/security.py", "禁止明文记录 token/api_key", "日志扫描"),
        ("权限分级", "mcp/security.py", "高危操作需确认", "权限回归"),
    ]
    for r in rows:
        c = t.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def d5_test_acceptance(d: Document) -> None:
    d.add_heading("本项目改动验收清单（开发自测+PR门禁）", 0)
    checks = [
        "受影响模块单元测试通过（tests/ 对应目录）。",
        "MCP 相关改动必须跑 tests/test_mcp.py 与 tests/test_mcp_input_validator.py。",
        "抓取或摘要改动需至少1条真实URL回归样例（可Mock外部依赖）。",
        "导出链路改动需验证 html/markdown/word 至少2种格式。",
        "lint 通过且不引入新的高优先级问题。",
    ]
    for i, c in enumerate(checks, 1):
        d.add_paragraph(f"[{i}] {c}")

    d.add_heading("发布闸门", 1)
    t = d.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "闸门"
    t.rows[0].cells[1].text = "阈值"
    t.rows[0].cells[2].text = "证据"
    for r in [
        ("功能", "关键路径100%通过", "回归报告"),
        ("安全", "高危问题0", "安全测试报告"),
        ("工程", "CI通过率>=95%", "流水线记录"),
    ]:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = r


def d6_handover(d: Document) -> None:
    d.add_heading("维护交接模板（仅针对本项目）", 0)
    d.add_paragraph("用于模块负责人交接，减少知识断层。")
    fields = [
        "模块名称：",
        "核心入口文件：",
        "关键依赖：",
        "已知风险：",
        "常见故障与处理：",
        "必须执行的回归用例：",
        "下阶段待办：",
    ]
    for f in fields:
        d.add_paragraph(f)


def main() -> None:
    docs = [
        ("01-项目开发维护总指南.docx", d1_arch_maint),
        ("02-模块级改进任务清单.docx", d2_module_tasks),
        ("03-本地开发运行手册.docx", d3_runbook),
        ("04-安全加固基线（项目定制）.docx", d4_security),
        ("05-改动验收清单（项目定制）.docx", d5_test_acceptance),
        ("06-维护交接模板（项目定制）.docx", d6_handover),
    ]
    paths = [mk(name, fn) for name, fn in docs]
    for p in paths:
        print(p)


if __name__ == "__main__":
    main()
